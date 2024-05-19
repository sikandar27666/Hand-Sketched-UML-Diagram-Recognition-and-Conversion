import pandas as pd
import numpy as np
from sklearn.model_selection import train_test_split, cross_val_score
from sklearn.metrics import confusion_matrix, classification_report, accuracy_score
from sklearn.linear_model import LogisticRegression
from sklearn.tree import DecisionTreeClassifier
from sklearn.neighbors import KNeighborsClassifier
from sklearn.naive_bayes import MultinomialNB
from sklearn.svm import SVC
from sklearn.ensemble import RandomForestClassifier
from sklearn.feature_extraction.text import CountVectorizer, TfidfVectorizer
import seaborn as sns
import matplotlib.pyplot as plt
from docx import Document

# Load the dataset
df = pd.read_csv('C:/Users/sk/Desktop/spam_Emails_data.csv')

# Check for NaN values in the text column and handle them
if df['text'].isna().sum() > 0:
    df['text'] = df['text'].fillna('')

# Assume 'text' is the column containing email text and 'label' is the target column ('spam' or 'not spam')
X = df['text']
y = df['label']

# Vectorize the text data using CountVectorizer or TfidfVectorizer
vectorizer = TfidfVectorizer(max_features=5000)
X = vectorizer.fit_transform(X)

# Split the data into training and testing sets
X_train, X_test, y_train, y_test = train_test_split(X, y, test_size=0.2, random_state=42)

# List of classifiers
classifiers = {
    'Logistic Regression': LogisticRegression(),
    'Decision Tree': DecisionTreeClassifier(),
    'k-Nearest Neighbors': KNeighborsClassifier(),
    'Naive Bayes': MultinomialNB(),
    'SVM': SVC(),
    'Random Forest': RandomForestClassifier()
}

# Dictionary to store the evaluation metrics and cross-validation results
results = {}

# Train and evaluate each classifier
for name, clf in classifiers.items():
    # Train the classifier
    clf.fit(X_train, y_train)

    # Predict on the test set
    y_pred = clf.predict(X_test)

    # Calculate accuracy, precision, recall, and F1-score
    accuracy = accuracy_score(y_test, y_pred)
    report = classification_report(y_test, y_pred, output_dict=True)

    # Store the evaluation metrics
    results[name] = {
        'Accuracy': accuracy,
        'Precision': report['weighted avg']['precision'],
        'Recall': report['weighted avg']['recall'],
        'F1-Score': report['weighted avg']['f1-score']
    }

    # Print confusion matrix
    print(f"Confusion Matrix for {name}:")
    conf_matrix = confusion_matrix(y_test, y_pred)
    sns.heatmap(conf_matrix, annot=True, fmt='d', cmap='Blues')
    plt.xlabel('Predicted')
    plt.ylabel('Actual')
    plt.title(f'Confusion Matrix for {name}')
    plt.show()

# Display the evaluation metrics table
results_df = pd.DataFrame(results).T
print(results_df)

# Perform k-fold cross-validation for each classifier
k = 5  # Number of folds
cv_results = {}

for name, clf in classifiers.items():
    scores = cross_val_score(clf, X, y, cv=k)
    cv_results[name] = {
        'Mean CV Accuracy': np.mean(scores),
        'Std CV Accuracy': np.std(scores)
    }
    print(f"{name} Cross-Validation Accuracy (mean): {np.mean(scores)}")
    print(f"{name} Cross-Validation Accuracy (std): {np.std(scores)}")

# Save results to a Word document

# Create a new Word document
doc = Document()

# Add a title
doc.add_heading('Classifier Performance on Spam Email Dataset', level=1)

# Add evaluation metrics table
doc.add_heading('Evaluation Metrics', level=2)
table = doc.add_table(rows=1, cols=5)
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Classifier'
hdr_cells[1].text = 'Accuracy'
hdr_cells[2].text = 'Precision'
hdr_cells[3].text = 'Recall'
hdr_cells[4].text = 'F1-Score'

for name, metrics in results.items():
    row = table.add_row().cells
    row[0].text = name
    row[1].text = f"{metrics['Accuracy']:.2f}"
    row[2].text = f"{metrics['Precision']:.2f}"
    row[3].text = f"{metrics['Recall']:.2f}"
    row[4].text = f"{metrics['F1-Score']:.2f}"

# Add cross-validation results
doc.add_heading('Cross-Validation Results', level=2)
cv_table = doc.add_table(rows=1, cols=3)
hdr_cells = cv_table.rows[0].cells
hdr_cells[0].text = 'Classifier'
hdr_cells[1].text = 'Mean CV Accuracy'
hdr_cells[2].text = 'Std CV Accuracy'

for name, metrics in cv_results.items():
    row = cv_table.add_row().cells
    row[0].text = name
    row[1].text = f"{metrics['Mean CV Accuracy']:.2f}"
    row[2].text = f"{metrics['Std CV Accuracy']:.2f}"

# Save the document
doc.save('classifier_results.docx')
